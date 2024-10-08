import io
import requests
import fitz
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
from docx import Document
from mistralai import Mistral
from bs4 import BeautifulSoup

NOMBRE_URL_PRIS = 10


import streamlit as st
APY_KEY = st.text_input("Clé API Mistral:")
CLE_GOOGLE = st.text_input("Clé Google:")
# Créer une liste déroulante (selectbox) avec différentes options
MODEL_LLM = st.selectbox(
    "Choix du modèle du moins couteux au plus performant:",
    ("open-mistral-7b", "mistral-small-2409", "mistral-large-2407")
)
option_model_Embedding = st.selectbox(
    "Modèle d'embedding du moins performant au plus performant:",
    ("paraphrase-MiniLM-L6-v2", "all-mpnet-base-v2", "sentence-t5-large")
)
   
if 'generated_files' not in st.session_state:
    st.session_state.generated_files = [] 
    
    
if 'text_inputs' not in st.session_state:
    st.session_state.text_inputs = []
if 'counter' not in st.session_state:
    st.session_state.counter = 0

# Fonction pour ajouter une nouvelle zone de texte
def add_text_input():
    st.session_state.counter += 1  # Augmente le compteur pour chaque zone ajoutée
    st.session_state.text_inputs.append("")

# Bouton pour ajouter une nouvelle zone de texte
if st.button("Ajouter une question:"):
    add_text_input()

# Afficher toutes les zones de texte
for i in range(st.session_state.counter):
    st.session_state.text_inputs[i] = st.text_input(f"Zone de texte {i+1}", st.session_state.text_inputs[i])
    
# Initialiser l'état de traitement dans session_state
if 'is_running' not in st.session_state:
    st.session_state.is_running = False

def get_google_search_results(mot_cle: str,num_results: int, cle_google=CLE_GOOGLE) -> list[str]:
    """Fonction qui renvoie des url à partir de mot clée en utilisant l'API google. Attention 100 requete par jour maximum.

    Args:
        mot_cle (str): _description_
        num_results (int): Nombre d'url voulu

    Returns:
        list[str]: liste d'url
    """    
    api_key = cle_google
    cx = "40a2133a6472a4f88"
    
    mot_cle = mot_cle.replace(" ", "+")  # Gérer les espaces dans la requête
    search_url = f"https://www.googleapis.com/customsearch/v1?q={mot_cle}&num={num_results}&key={api_key}&cx={cx}"
    response = requests.get(search_url)
    response.raise_for_status()
    results = response.json()
    links = []
    if 'items' in results:
        for item in results['items']:
            links.append(item['link'])
            if len(links) >= num_results:
                break
    return links


def transformer(sentence1:str, MODEL)->float:
    """Compare la cos similarité entre deux phrases

    Args:
        sentence1 (str): phrase 1
        sentence2 (str): phrase 2

    Returns:
        float: la cosimilarité
    """    
    # On utilise un modèle SBert assez léger car j'ai un caillou
    
    embedding1 = MODEL.encode(sentence1)

    return embedding1

def compare_deux_trans(vect1: float,vect2: float)->float:
    similarity = cosine_similarity([vect1], [vect2])
    return similarity[0][0]




def extract_text_from_web(url: str) -> list[str]:
    """
    Prend une URL et renvoie le contenu des balises <p> de la page web sous forme de liste.

    Args:
        url (str): L'URL de la page web à analyser.

    Returns:
        list[str]: Une liste contenant le texte de chaque balise <p> de la page web.
    """
    response = requests.get(url)
    soup = BeautifulSoup(response.content, 'html.parser')
    
    return soup.get_text()


def extract_text_from_pdf(url: str) -> list[str]:
    response = requests.get(url)
    with open('temp.pdf', 'wb') as f:
        f.write(response.content)
    doc = fitz.open('temp.pdf')
    text = ''
    for page in doc:
        text += page.get_text()
    return text

# Fonction pour diviser le texte en chunks
def chunk_text(text: str, min_chunk_size: int=500, max_chunk_size: int=800) -> list[str]:
    chunks = []
    
    if len(text) > max_chunk_size:
        paragraphs = text.split("\n")
        for para in paragraphs:
            if len(para) > max_chunk_size:
                words = para.split()
                chunk = []
                chunk_length = 0
                for word in words:
                    if chunk_length + len(word) + 1 <= max_chunk_size:
                        chunk.append(word)
                        chunk_length += len(word) + 1
                    else:
                        if chunk_length >= min_chunk_size:
                            chunks.append(" ".join(chunk))
                            chunk = [word]
                            chunk_length = len(word) + 1
                        else:
                            chunk.append(word)
                            chunk_length += len(word) + 1
                if chunk:
                    chunks.append(" ".join(chunk))
            else:
                chunks.append(para)
    else:
        chunks.append(text)

    # Ensure all chunks meet the minimum size requirement
    final_chunks = []
    temp_chunk = ""
    for chunk in chunks:
        if len(temp_chunk) + len(chunk) + 1 <= max_chunk_size:
            temp_chunk += chunk + " "
        else:
            if len(temp_chunk) >= min_chunk_size:
                final_chunks.append(temp_chunk.strip())
                temp_chunk = chunk + " "
            else:
                temp_chunk += chunk + " "
    if temp_chunk.strip():
        final_chunks.append(temp_chunk.strip())

    return final_chunks


# Fonction principale pour traiter une URL
def process_url(url: str) -> list[str]:
    if url.endswith('.pdf'):
        text = extract_text_from_pdf(url)
    else:
        text = extract_text_from_web(url)
    chunks = chunk_text(text)
    return chunks


def write_prompt(prompt: str, api_key = APY_KEY, model: str = MODEL_LLM) -> str:

    client = Mistral(api_key=api_key)

    chat_response = client.chat.complete(
        model= model,
        messages = [
            {
                "role": "user",
                "content": prompt,
            },
        ]
    )
    return chat_response.choices[0].message.content

def reformuler_recherche(liste_question: str, api_key: str = APY_KEY) -> str:
    liste_question_reformuler=[]
    for question in liste_question:
        prompt_cos=f"""Je veux que tu répondes strictement à la question en reprenant le plus de fois possible les mots important de la question après les deux points :{question}."""
        liste_question_reformuler.append(write_prompt(prompt_cos, api_key))
    return liste_question_reformuler


 
def creation_liste_question_vect_question(liste_question: str, MODEL) -> list[list[float,str]]:
    """transforme la liste de question en vecteur

    Args:
        liste_question (str): liste des questions de l'utilisateur

    Returns:
        list[list[float,str]]: une liste de liste qui contient le vecteur de la question et la question
    """    
    liste_question_vect_question = [[transformer(question, MODEL), question] for question in liste_question]
    return liste_question_vect_question
        

def r_dico_chunk_poid(nombre_chunk_pris: int, liste_question: str) -> dict[list[str, float]]:
    """Ca mets dans un dico avec comme clés la questions et les chunks avec la valeur de 0
    Args:
        nombre_chunk_pris (int): Le nombre de chunk que l'on va mettre dans le LLM
        liste_question (str): liste des questions de l'utilisateur

    Returns:
        dict[list[str, float]]: dictionnaire d'une liste qui contient le chunk et 0
    """    
    dico_chunk_poid = dict()
    for i in liste_question:
        dico_chunk_poid[i] = []
        for _ in range(nombre_chunk_pris):        
            dico_chunk_poid[i].append(["",0,""])
    return dico_chunk_poid
    

def trouver_list_url(liste_question: str, NOMBRE_URL_PRIS: int) -> list[str]:
    """Renvoies la liste des urls 

    Args:
        liste_question (str): _description_
        NOMBRE_URL_PRIS (int): _description_

    Returns:
        list[str]: _description_
    """    
    liste_url = []
    for i in liste_question:
        liste_url+=get_google_search_results(i, NOMBRE_URL_PRIS)
    liste_url=set(liste_url)
    return liste_url


def trouver_meilleur_chunk(dico_chunk_poid: dict, chunks: list, liste_question_vect_question: list, nombre_pris: int, MODEL) -> dict:
    for i,z in enumerate(chunks):
        if z and len(z) > 0:
            for phrase in z[0]:
                for j,k in liste_question_vect_question: 
                    chunk_vect=transformer(phrase, MODEL)
                    cos=compare_deux_trans(chunk_vect,j)
                    for m in range(nombre_pris):
                        if(dico_chunk_poid[k][-1*(m)-1][1]<cos):
                                dico_chunk_poid[k][-1*m-1][0]=phrase
                                dico_chunk_poid[k][-1*m-1][1]=cos
                                dico_chunk_poid[k][-1*m-1][2]=z[1]
                                break
        else:
            continue
    return dico_chunk_poid

def creation_document(dico_chunk_poid: dict, nombre_pris: int, liste_question: list[str]):
        
    for nombre_question,i in enumerate(dico_chunk_poid.keys()): 
        document = Document()
        base_donne=""
        source=""
        for j in range(nombre_pris):
            base_donne+=dico_chunk_poid[i][j][0]
            base_donne+="\n\n" 
            source+=dico_chunk_poid[i][j][2] +"\n\n\n"
            source+=dico_chunk_poid[i][j][0] + "\n\n\n\n\n"
        prompt=f"""A partir des données réponds à {i},\nVoici les données :{base_donne}\nréponds à la question {i} en détaillant  avec les données fournis plus haut"""
        reponse = write_prompt(prompt, APY_KEY)
        document.add_heading(liste_question[nombre_question], level=1)
        document.add_paragraph(reponse)
        st.write(reponse)
        document.add_heading("Source", level=1)
        st.title(reponse)
        document.add_paragraph(source)
        st.write(source)
        # Sauvegarde en mémoire (buffer)
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        # Ajouter le fichier généré à la liste des fichiers dans session_state
        st.session_state.generated_files.append({
            "file_name": f"reponse_question_{nombre_question+1}.docx",
            "data": buffer
        })
         
def pipeline_rag_internet(nombre_pris: int = 5):
    MODEL = SentenceTransformer(option_model_Embedding)
    liste_question = st.session_state.text_inputs
    reformuler = reformuler_recherche(liste_question)
    liste_question_vect_question = creation_liste_question_vect_question(reformuler, MODEL)
    dico_chunk_poid = r_dico_chunk_poid(nombre_pris, reformuler)
    liste_url = trouver_list_url(liste_question, NOMBRE_URL_PRIS)
    chunks=[]
    for url in liste_url:
        print(f"url:{url}")
        chunk=process_url(url)
        chunks.append([chunk,url])
    dico_chunk_poid = trouver_meilleur_chunk(dico_chunk_poid, chunks, liste_question_vect_question, nombre_pris, MODEL)
    creation_document(dico_chunk_poid, nombre_pris, liste_question)



# Définir une fonction longue à exécuter
def longue_fonction():
    pipeline_rag_internet()

# Gérer le clic sur le bouton
if st.button("Exécuter la fonction", disabled=st.session_state.is_running):
    st.session_state.is_running = True  # Désactiver le bouton immédiatement
    longue_fonction()  # Exécuter la fonction longue
    st.session_state.is_running = False  # Réactiver le bouton après exécution

# Afficher un message tant que la fonction est en cours d'exécution
if st.session_state.is_running:
    st.write("Veuillez attendre, la fonction est en cours d'exécution...")

for file_info in st.session_state.generated_files:
    st.download_button(
        label=f"Télécharger {file_info['file_name']}",
        data=file_info['data'],
        file_name=file_info['file_name'],
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )